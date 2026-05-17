"""
UrbanRenewal Planner — Streamlit Web 界面

运行方式：
    streamlit run app.py

额外依赖（除项目已有包外）：
    pip install streamlit-folium fpdf2 python-docx
"""

from __future__ import annotations

import io
import sys
import time
from pathlib import Path
from typing import Optional

# ---------- 路径 ----------
sys.path.insert(0, str(Path(__file__).parent))

import folium
import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from streamlit_folium import st_folium

load_dotenv(Path(__file__).parent / ".env", override=False)

from src.urbanrenewal.agent.planner import build_graph, PlannerState
from src.urbanrenewal.config import cfg
from src.urbanrenewal.tools.poi_query import query_poi_by_buffer
from src.urbanrenewal.tools.policy_rag import format_citations

# ============================================================
# 1. 页面配置
# ============================================================

st.set_page_config(
    page_title="UrbanRenewal Planner · 杨浦区城市更新",
    page_icon="🏙️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------- 全局 CSS ----------
st.markdown("""
<style>
/* ---- 全局背景与默认文字 ---- */
[data-testid="stAppViewContainer"] {
    background: #0d1117;
}
.stApp {
    color: #c9d1d9;
}

/* ---- 侧边栏 ---- */
[data-testid="stSidebar"] {
    background: #161b22;
    border-right: 1px solid #30363d;
}
/* 覆盖侧边栏所有文字颜色 */
[data-testid="stSidebar"] * {
    color: #c9d1d9 !important;
}
[data-testid="stSidebar"] .stMarkdown p,
[data-testid="stSidebar"] .stMarkdown span {
    color: #c9d1d9 !important;
}
/* caption 单独提亮 */
[data-testid="stSidebar"] [data-testid="stCaptionContainer"] p {
    color: #8b949e !important;
    font-size: 0.82rem !important;
}
/* 侧边栏 divider */
[data-testid="stSidebar"] hr {
    border-color: #30363d !important;
}

/* ---- 标题区 ---- */
.app-header {
    padding: 0.5rem 0 0.8rem 0;
    border-bottom: 1px solid #21262d;
    margin-bottom: 1rem;
}
.app-title {
    font-size: 1.75rem;
    font-weight: 700;
    color: #58a6ff;
    letter-spacing: 0.01em;
    margin: 0;
    line-height: 1.2;
}
.app-subtitle {
    font-size: 0.9rem;
    color: #8b949e;
    margin-top: 0.3rem;
}

/* ---- 侧边栏分区标题 ---- */
.sidebar-section {
    font-size: 0.72rem;
    font-weight: 600;
    color: #58a6ff !important;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    margin: 1.1rem 0 0.4rem 0;
    padding-bottom: 0.25rem;
    border-bottom: 1px solid #21262d;
}

/* ---- 模型信息条 ---- */
.model-badge {
    display: inline-block;
    background: #21262d;
    border: 1px solid #30363d;
    border-radius: 5px;
    padding: 3px 8px;
    font-size: 0.78rem;
    color: #79c0ff !important;
    margin: 2px 0;
    font-family: monospace;
}

/* ---- 聊天气泡 ---- */
.msg-user {
    background: #1c2d40;
    border-left: 3px solid #58a6ff;
    border-radius: 0 8px 8px 0;
    padding: 0.75rem 1.1rem;
    margin: 0.5rem 0;
    color: #cae8ff;
    font-size: 0.95rem;
}
.msg-agent {
    background: #172011;
    border-left: 3px solid #3fb950;
    border-radius: 0 8px 8px 0;
    padding: 0.75rem 1.1rem;
    margin: 0.5rem 0;
    color: #aff5b4;
    font-size: 0.95rem;
}
.msg-error {
    background: #2d1113;
    border-left: 3px solid #f85149;
    border-radius: 0 8px 8px 0;
    padding: 0.75rem 1.1rem;
    margin: 0.5rem 0;
    color: #ffa198;
    font-size: 0.95rem;
}

/* ---- 指标卡片 ---- */
.metric-row {
    display: flex;
    gap: 0.75rem;
    margin: 0.75rem 0 1.25rem 0;
}
.metric-card {
    flex: 1;
    background: #161b22;
    border: 1px solid #30363d;
    border-radius: 10px;
    padding: 0.9rem 0.8rem;
    text-align: center;
}
.metric-card:hover {
    border-color: #58a6ff;
    transition: border-color 0.2s;
}
.metric-value {
    font-size: 1.9rem;
    font-weight: 700;
    color: #58a6ff;
    line-height: 1.1;
}
.metric-label {
    font-size: 0.68rem;
    color: #8b949e;
    text-transform: uppercase;
    letter-spacing: 0.07em;
    margin-top: 0.2rem;
}

/* ---- 设施缺口标签 ---- */
.gap-wrap { display: flex; flex-wrap: wrap; gap: 4px; margin-top: 4px; }
.tag-missing {
    background: #3d1f1f; color: #ff7b72;
    border: 1px solid #5c2424;
    border-radius: 5px; padding: 3px 9px; font-size: 0.78rem;
}
.tag-low {
    background: #2d2008; color: #e3b341;
    border: 1px solid #4a3510;
    border-radius: 5px; padding: 3px 9px; font-size: 0.78rem;
}
.tag-ok {
    background: #0f2d17; color: #56d364;
    border: 1px solid #1a4028;
    border-radius: 5px; padding: 3px 9px; font-size: 0.78rem;
}

/* ---- 示例问题按钮 ---- */
.example-grid { display: flex; flex-wrap: wrap; gap: 0.5rem; margin: 0.5rem 0 1.5rem 0; }
.example-btn {
    background: #161b22;
    border: 1px solid #30363d;
    border-radius: 6px;
    padding: 0.45rem 0.9rem;
    color: #79c0ff;
    font-size: 0.85rem;
    cursor: pointer;
    transition: background 0.15s, border-color 0.15s;
}
.example-btn:hover { background: #1c2d40; border-color: #58a6ff; }

/* ---- Streamlit 默认组件微调 ---- */
/* 主内容文字 */
.stMarkdown p, .stMarkdown li { color: #c9d1d9; }
.stMarkdown h1, .stMarkdown h2, .stMarkdown h3 { color: #e6edf3; }
/* subheader */
h3[data-testid] { color: #e6edf3; }
/* expander */
[data-testid="stExpander"] summary { color: #79c0ff; }
/* dataframe */
[data-testid="stDataFrame"] { border: 1px solid #30363d; border-radius: 8px; overflow: hidden; }
/* progress bar */
[data-testid="stProgressBar"] > div > div { background: #58a6ff; }
/* divider */
hr { border-color: #21262d !important; }

/* 隐藏默认水印 */
#MainMenu, footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)


# ============================================================
# 2. Session State
# ============================================================

def _init_state() -> None:
    defaults: dict = {
        "messages": [],       # list[{role, content, state?, is_error?}]
        "last_result": None,  # 最近一次 Agent 完整 state dict
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init_state()


# ============================================================
# 3. 导出工具函数
# ============================================================

def _find_chinese_font() -> Optional[str]:
    """查找系统可用的中文 TTF 字体路径（Windows）。"""
    candidates = [
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simfang.ttf",
        "C:/Windows/Fonts/SimsunExtG.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return p
    return None


def _build_docx(state: dict) -> bytes:
    """将 Agent 结果导出为 Word 文档，返回 bytes。"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    doc = Document()
    doc.core_properties.author = "UrbanRenewal Planner Agent"

    h = doc.add_heading("城市更新规划诊断报告", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    geo = state.get("geocode_result")
    if geo:
        doc.add_paragraph(f"分析地点：{geo.address}")
        doc.add_paragraph(f"坐标：WGS84({geo.lon_wgs84:.5f}, {geo.lat_wgs84:.5f})")

    intent = state.get("intent") or {}
    doc.add_paragraph(
        f"分析场景：{intent.get('scenario', '—')}  "
        f"分析半径：{intent.get('radius_m', 800)}m"
    )

    doc.add_heading("规划建议", level=1)
    doc.add_paragraph(state.get("answer", ""))

    poi_summary = state.get("poi_summary") or []
    if poi_summary:
        doc.add_heading("设施统计", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light List Accent 1"
        for cell, hdr in zip(
            tbl.rows[0].cells,
            ["类别", "数量", "平均距离(m)", "最近(m)"],
        ):
            cell.text = hdr
        for row in poi_summary:
            cells = tbl.add_row().cells
            cells[0].text = str(row.get("category_planning", ""))
            cells[1].text = str(row.get("count", ""))
            cells[2].text = str(row.get("avg_distance_m", ""))
            cells[3].text = str(row.get("min_distance_m", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(state: dict) -> bytes:
    """
    将 Agent 结果导出为 PDF，返回 bytes。

    优先加载 Windows 系统中文字体（SimHei 等）支持中文渲染；
    若找不到则仅输出英文 ASCII 内容。
    """
    try:
        from fpdf import FPDF
    except ImportError:
        return b""

    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(15, 15, 15)

    font_path = _find_chinese_font()
    has_chinese_font = font_path is not None

    if has_chinese_font:
        pdf.add_font("CJK", fname=font_path)
        pdf.set_font("CJK", size=14)
    else:
        pdf.set_font("Helvetica", size=14)

    # 标题
    geo = state.get("geocode_result")
    title = f"城市更新规划报告 — {geo.address}" if geo else "城市更新规划报告"
    if not has_chinese_font:
        title = "Urban Renewal Planning Report"

    pdf.cell(0, 10, title, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    if has_chinese_font:
        pdf.set_font("CJK", size=10)
    else:
        pdf.set_font("Helvetica", size=10)

    # 地点信息
    if geo:
        pdf.set_font_size(11)
        pdf.cell(0, 8,
                 f"地点：{geo.address}  坐标：({geo.lon_wgs84:.5f}, {geo.lat_wgs84:.5f})",
                 new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # 规划建议正文
    answer = state.get("answer", "")
    if has_chinese_font:
        pdf.set_font("CJK", size=10)
        pdf.multi_cell(0, 6, answer)
    else:
        pdf.set_font("Helvetica", size=10)
        safe = answer.encode("ascii", errors="replace").decode()
        pdf.multi_cell(0, 6, safe)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ============================================================
# 4. 侧边栏
# ============================================================

with st.sidebar:
    st.markdown(
        '<p style="font-size:1.3rem;font-weight:700;color:#58a6ff;margin:0.4rem 0 0 0">'
        "🏙️ UrbanRenewal</p>",
        unsafe_allow_html=True,
    )
    st.markdown(
        '<p style="font-size:0.82rem;color:#8b949e;margin:0.1rem 0 0.6rem 0">'
        "杨浦区城市更新规划 Agent</p>",
        unsafe_allow_html=True,
    )
    st.divider()

    # —— 模型信息 ——
    st.markdown('<p class="sidebar-section">运行环境</p>', unsafe_allow_html=True)
    for label, value in [
        ("LLM", cfg.llm_model),
        ("VL", cfg.vl_model),
        ("Embed", cfg.rag_embedding_model),
    ]:
        st.markdown(
            f'<div><span style="color:#8b949e;font-size:.78rem">{label}：</span>'
            f'<span class="model-badge">{value}</span></div>',
            unsafe_allow_html=True,
        )

    st.divider()

    # —— 操作按钮 ——
    if st.button("🗑️ 清空对话", use_container_width=True):
        st.session_state.messages = []
        st.session_state.last_result = None
        st.rerun()

    # —— 导出（有结果时才显示）——
    last = st.session_state.last_result
    if last and last.get("answer"):
        st.divider()
        st.markdown('<p class="sidebar-section">导出结果</p>', unsafe_allow_html=True)
        col_a, col_b = st.columns(2)
        with col_a:
            st.download_button(
                "📄 Word",
                data=_build_docx(last),
                file_name="规划建议.docx",
                mime="application/vnd.openxmlformats-officedocument"
                      ".wordprocessingml.document",
                use_container_width=True,
            )
        with col_b:
            st.download_button(
                "📑 PDF",
                data=_build_pdf(last),
                file_name="规划建议.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

    st.divider()

    # —— 使用说明 ——
    st.markdown('<p class="sidebar-section">使用说明</p>', unsafe_allow_html=True)
    st.markdown(
        '<p style="font-size:0.8rem;color:#8b949e;line-height:1.6">'
        "直接用自然语言描述规划问题，Agent 会自动识别地点、"
        "场景和分析范围，并综合 POI、路网、街景与政策文献"
        "生成结构化建议。<br><br>"
        "当前仅支持<b style='color:#58a6ff'>上海市杨浦区</b>范围内的问题。"
        "</p>",
        unsafe_allow_html=True,
    )


# ============================================================
# 5. Agent 工作流（带进度条）
# ============================================================

def _run_agent_with_progress(question: str) -> dict:
    """
    运行 LangGraph Agent，通过 stream 逐节点更新进度条。
    Agent 完全自主判断场景、半径和是否需要街景分析。
    返回最终合并后的 state dict。
    """
    progress_bar = st.progress(0, text="正在启动分析…")
    status_placeholder = st.empty()

    def _step(pct: int, msg: str) -> None:
        progress_bar.progress(pct, text=msg)
        status_placeholder.markdown(
            f'<p style="color:#8b949e;font-size:0.82rem;margin:0">{msg}</p>',
            unsafe_allow_html=True,
        )
        time.sleep(0.04)

    _step(5, "🚀 正在启动 Agent…")

    graph = build_graph()
    # 完全依赖 parse_intent 节点自主识别意图，不预设任何参数
    initial: PlannerState = {
        "question": question,
        "intent": {},
        "geocode_result": None,
        "poi_summary": None,
        "gap_report": None,
        "isochrone": None,
        "intersections": [],
        "sv_results": [],
        "policy_chunks": [],
        "answer": "",
        "error": "",
    }

    NODE_PROGRESS: dict[str, tuple[int, str]] = {
        "parse_intent":     (12, "🧠 解析问题意图与空间范围…"),
        "geocode":          (26, "📍 地理编码，获取精确坐标…"),
        "query_spatial":    (48, "🗺️ 检索周边 POI 与路网数据…"),
        "query_streetview": (64, "📷 分析街景图片…"),
        "query_policy":     (78, "📚 检索城市更新政策文献…"),
        "generate":         (92, "✍️ 综合数据，生成规划建议…"),
    }

    final_state: dict = dict(initial)

    try:
        for event in graph.stream(initial, stream_mode="updates"):
            node_name = next(iter(event))
            node_out: dict = event[node_name]
            for k, v in node_out.items():
                if v is not None:
                    final_state[k] = v
            if node_name in NODE_PROGRESS:
                pct, msg = NODE_PROGRESS[node_name]
                _step(pct, msg)

        _step(100, "✅ 分析完成")
    except Exception as exc:
        progress_bar.empty()
        status_placeholder.empty()
        return {**initial, "error": str(exc),
                "answer": f"Agent 运行出错：{exc}"}

    progress_bar.empty()
    status_placeholder.empty()
    return final_state


# ============================================================
# 6. 地图渲染
# ============================================================

def _render_map(state: dict) -> Optional[folium.Map]:
    """根据 Agent state 渲染多图层 Folium 地图。"""
    geo = state.get("geocode_result")
    if not geo:
        return None

    lon, lat = geo.lon_wgs84, geo.lat_wgs84
    radius_m = (state.get("intent") or {}).get("radius_m", 800)

    m = folium.Map(
        location=[lat, lon],
        zoom_start=15,
        tiles="CartoDB dark_matter",
        prefer_canvas=True,
    )

    # 分析中心点
    folium.Marker(
        [lat, lon],
        popup=folium.Popup(geo.address, max_width=240),
        tooltip=f"📍 {geo.address}",
        icon=folium.Icon(color="blue", icon="map-marker", prefix="fa"),
    ).add_to(m)

    # 分析半径圆
    folium.Circle(
        [lat, lon],
        radius=radius_m,
        color="#58a6ff",
        weight=1.5,
        fill=True,
        fill_opacity=0.05,
        tooltip=f"分析范围 {radius_m}m",
    ).add_to(m)

    # 步行等时圈凸包
    iso = state.get("isochrone")
    if iso and iso.hull_polygon:
        coords = [[c[1], c[0]] for c in iso.hull_polygon.exterior.coords]
        folium.Polygon(
            coords,
            color="#3fb950",
            weight=1.5,
            fill=True,
            fill_opacity=0.07,
            tooltip="步行 10 分钟可达范围",
        ).add_to(m)

    # POI 点位（按类别着色）
    sv_scenario = (state.get("intent") or {}).get("scenario", "general")
    try:
        poi_gdf = query_poi_by_buffer(
            lon, lat,
            radius_m=radius_m,
            scenario=sv_scenario if sv_scenario != "general" else None,
        )
    except Exception:
        poi_gdf = None

    CAT_COLORS = {
        "医疗": "red",       "药店": "pink",       "养老服务": "orange",
        "公园绿地": "green", "公交站": "blue",      "地铁站": "darkblue",
        "菜场": "cadetblue", "社区服务设施": "purple", "公厕": "lightgray",
    }
    if poi_gdf is not None and not poi_gdf.empty:
        poi_layer = folium.FeatureGroup(name="🏪 周边设施", show=True)
        for _, row in poi_gdf.head(250).iterrows():
            cat = str(row.get("category_planning", ""))
            folium.CircleMarker(
                [row["lat_wgs84"], row["lon_wgs84"]],
                radius=4,
                color=CAT_COLORS.get(cat, "gray"),
                fill=True,
                fill_opacity=0.8,
                tooltip=f"{row['name']}  {cat}  {row['distance_m']}m",
            ).add_to(poi_layer)
        poi_layer.add_to(m)

    # 街景分析点位
    sv_results = state.get("sv_results") or []
    if sv_results:
        sv_layer = folium.FeatureGroup(name="📷 街景分析", show=True)
        for r in sv_results:
            preview = (r.analysis_text or "")[:120]
            folium.Marker(
                [r.meta.lat_wgs84, r.meta.lon_wgs84],
                popup=folium.Popup(
                    f"<b>{r.meta.image_id}</b><br>"
                    f"朝向：{r.meta.direction_bucket} · {r.meta.distance_m}m<br>"
                    f"<small>{preview}…</small>",
                    max_width=300,
                ),
                tooltip=f"街景 {r.meta.direction_bucket}向 {r.meta.distance_m}m",
                icon=folium.Icon(color="purple", icon="camera", prefix="fa"),
            ).add_to(sv_layer)
        sv_layer.add_to(m)

    # 重要路口
    intersections = state.get("intersections") or []
    if intersections:
        int_layer = folium.FeatureGroup(name="🔀 重要路口", show=False)
        for r in intersections[:20]:
            folium.CircleMarker(
                [r.lat, r.lon],
                radius=5,
                color="#e3b341",
                fill=True,
                fill_opacity=0.85,
                tooltip=f"路口 {r.street_count}路 {r.distance_m}m",
            ).add_to(int_layer)
        int_layer.add_to(m)

    folium.LayerControl(collapsed=False, position="topright").add_to(m)
    return m


# ============================================================
# 7. 结果展示区
# ============================================================

def _render_results(state: dict) -> None:
    """渲染完整结果：指标卡 + 建议文本 + 地图 + 折叠详情。"""
    geo = state.get("geocode_result")
    if not geo:
        if state.get("answer"):
            st.warning(state["answer"])
        return

    # —— 四个指标卡 ——
    poi_summary = state.get("poi_summary") or []
    iso = state.get("isochrone")
    intersections = state.get("intersections") or []
    policy_chunks = state.get("policy_chunks") or []
    total_poi = sum(r.get("count", 0) for r in poi_summary)

    cards_html = '<div class="metric-row">'
    for val, lbl in [
        (total_poi,                      "范围内设施"),
        (iso.node_count if iso else "—", "步行10分钟节点"),
        (len(intersections),             "周边路口"),
        (len(policy_chunks),             "政策依据"),
    ]:
        cards_html += (
            f'<div class="metric-card">'
            f'<div class="metric-value">{val}</div>'
            f'<div class="metric-label">{lbl}</div>'
            f'</div>'
        )
    cards_html += "</div>"
    st.markdown(cards_html, unsafe_allow_html=True)

    # —— 地址提示条 ——
    intent = state.get("intent") or {}
    in_yangpu = geo.in_target_district
    loc_color = "#3fb950" if in_yangpu else "#e3b341"
    loc_label = "已确认在杨浦区" if in_yangpu else "行政区信息不确定，以坐标为准"
    st.markdown(
        f'<p style="font-size:0.85rem;color:#8b949e;margin:0 0 0.8rem 0">'
        f'📍 <span style="color:#c9d1d9">{geo.address}</span>'
        f'&nbsp;&nbsp;<span style="color:{loc_color};font-size:0.78rem">'
        f'[{loc_label}]</span>'
        f'&nbsp;&nbsp;场景：<span style="color:#79c0ff">{intent.get("scenario","—")}</span>'
        f'&nbsp;·&nbsp;半径：<span style="color:#79c0ff">{intent.get("radius_m",800)}m</span>'
        f'</p>',
        unsafe_allow_html=True,
    )

    # —— 左：建议+诊断 / 右：地图 ——
    left, right = st.columns([1, 1], gap="large")

    with left:
        answer = state.get("answer", "")
        if answer:
            st.markdown(
                '<h3 style="color:#e6edf3;margin-bottom:.5rem">📋 规划建议</h3>',
                unsafe_allow_html=True,
            )
            st.markdown(answer)

        gap = state.get("gap_report")
        if gap:
            st.markdown(
                '<h3 style="color:#e6edf3;margin:1.2rem 0 .5rem 0">📊 设施配置诊断</h3>',
                unsafe_allow_html=True,
            )
            g1, g2, g3 = st.columns(3)
            for col, items, title, css in [
                (g1, gap.missing,      "完全缺失", "tag-missing"),
                (g2, gap.insufficient, "数量不足", "tag-low"),
                (g3, gap.present,      "已配置",   "tag-ok"),
            ]:
                with col:
                    st.markdown(
                        f'<p style="font-size:.8rem;color:#8b949e;'
                        f'margin:0 0 4px 0;font-weight:600">{title}</p>',
                        unsafe_allow_html=True,
                    )
                    if items:
                        tags = "".join(
                            f'<span class="{css}">{it}</span> ' for it in items
                        )
                        st.markdown(
                            f'<div class="gap-wrap">{tags}</div>',
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown(
                            '<span style="color:#3fb950;font-size:.8rem">无缺口</span>',
                            unsafe_allow_html=True,
                        )

        if policy_chunks:
            with st.expander("📚 政策依据原文", expanded=False):
                st.markdown(
                    format_citations(
                        policy_chunks, max_text_length=200, group_by_doc=True
                    )
                )

    with right:
        st.markdown(
            '<h3 style="color:#e6edf3;margin-bottom:.5rem">🗺️ 空间分析地图</h3>',
            unsafe_allow_html=True,
        )
        fmap = _render_map(state)
        if fmap:
            st_folium(fmap, width=None, height=500, returned_objects=[])
        else:
            st.info("地图数据加载中…")

    # —— 折叠详情 ——
    if poi_summary:
        with st.expander(
            f"🏪 周边设施详细统计（{len(poi_summary)} 类）", expanded=False
        ):
            df = pd.DataFrame(poi_summary)
            df.columns = ["设施类别", "数量", "平均距离(m)", "最近距离(m)"]
            st.dataframe(
                df.style.background_gradient(subset=["数量"], cmap="Blues").format(
                    {"数量": "{:,}", "平均距离(m)": "{:.0f}", "最近距离(m)": "{:.0f}"}
                ),
                use_container_width=True,
                hide_index=True,
            )

    sv_results = state.get("sv_results") or []
    if sv_results:
        with st.expander(
            f"📷 街景分析结果（{len(sv_results)} 张）", expanded=False
        ):
            for r in sv_results:
                tag_color = "#8b949e" if r.from_cache else "#3fb950"
                tag_text = "缓存命中" if r.from_cache else "新分析"
                st.markdown(
                    f'<p style="margin:.4rem 0 .2rem 0">'
                    f'<b style="color:#c9d1d9">{r.meta.image_id}</b>'
                    f'&nbsp;·&nbsp;{r.meta.direction_bucket}向'
                    f'&nbsp;·&nbsp;{r.meta.distance_m}m'
                    f'&nbsp;<span style="color:{tag_color};font-size:.75rem">'
                    f'[{tag_text}]</span></p>',
                    unsafe_allow_html=True,
                )
                if r.analysis_text and not r.analysis_text.startswith("["):
                    st.markdown(r.analysis_text)
                st.divider()


# ============================================================
# 8. 消息历史渲染
# ============================================================

def _render_messages() -> None:
    """渲染历史消息气泡及对应结果区域。"""
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(
                f'<div class="msg-user">🧑‍💼&nbsp; {msg["content"]}</div>',
                unsafe_allow_html=True,
            )
        elif msg["role"] == "assistant":
            if msg.get("is_error"):
                st.markdown(
                    f'<div class="msg-error">⚠️&nbsp; {msg["content"]}</div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    '<div class="msg-agent">'
                    "🤖&nbsp; 规划建议已生成，详见下方结果区域。"
                    "</div>",
                    unsafe_allow_html=True,
                )
                if msg.get("state"):
                    _render_results(msg["state"])


# ============================================================
# 9. 主界面
# ============================================================

# 顶部标题区
st.markdown(
    '<div class="app-header">'
    '<h1 class="app-title">🏙️ UrbanRenewal Planner</h1>'
    '<p class="app-subtitle">'
    "上海市杨浦区城市更新规划诊断与建议生成 Agent &nbsp;·&nbsp; "
    "输入自然语言问题，Agent 自动完成地理编码 → 空间分析 → 政策检索 → 建议生成"
    "</p>"
    "</div>",
    unsafe_allow_html=True,
)

# 历史消息
_render_messages()

# 聊天输入框
question_input: Optional[str] = st.chat_input(
    placeholder="请输入规划问题，例如：请分析鞍山新村周边800米的老年友好问题"
)

# 首次进入时显示示例问题
if not st.session_state.messages:
    st.markdown(
        '<p style="color:#8b949e;font-size:.85rem;margin-bottom:.4rem">'
        "💡 快速开始 — 点击下方示例问题"
        "</p>",
        unsafe_allow_html=True,
    )
    EXAMPLES = [
        ("🧓", "请分析鞍山新村周边800米的老年友好问题"),
        ("🚶", "控江路和本溪路路口有哪些步行环境问题？"),
        ("🏘️", "杨浦区平凉路社区15分钟生活圈设施是否完善？"),
        ("🏥", "请分析新华医院周边的适老化改造需求"),
    ]
    c1, c2 = st.columns(2)
    for i, (icon, ex) in enumerate(EXAMPLES):
        with (c1 if i % 2 == 0 else c2):
            if st.button(f"{icon} {ex}", key=f"ex_{i}", use_container_width=True):
                question_input = ex

# 处理新输入
if question_input:
    st.session_state.messages.append({"role": "user", "content": question_input})
    st.rerun()

# 有待处理的 user 消息 → 触发 Agent
messages = st.session_state.messages
if messages and messages[-1]["role"] == "user":
    question = messages[-1]["content"]

    result = _run_agent_with_progress(question=question)

    st.session_state.last_result = result
    answer = result.get("answer", "")
    error = result.get("error", "")

    if error and error != "out_of_scope" and not answer:
        st.session_state.messages.append({
            "role": "assistant",
            "content": f"分析失败：{error}",
            "is_error": True,
        })
    else:
        st.session_state.messages.append({
            "role": "assistant",
            "content": answer,
            "state": result,
            "is_error": False,
        })

    st.rerun()
